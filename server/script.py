from flask import Flask, request, render_template
from flask_cors import CORS
from docx import Document

app = Flask(__name__)
CORS(app)

# @app.route('/', methods=['GET', 'POST'])
# def hello():
#     if request.method == 'POST':
#         file = request.files['uploadedFile']
#         doc = Document(file)
#         new_doc = []
#         for paragraph in doc.paragraphs:
#             new_doc.append(paragraph.text)
        
#         if new_doc:
#             print(new_doc) 
    
#     return 'hello'


@app.route('/', methods=['GET', 'POST'])
def hello():
    new_doc = None

    if request.method == 'POST':
        file = request.files['uploadedFile']
        doc = Document(file)
        new_doc = []
        for paragraph in doc.paragraphs:
            new_doc.append(paragraph.text)

    return render_template('index.html', new_doc=new_doc)

if __name__ == '__main__':
    app.run(debug=True)